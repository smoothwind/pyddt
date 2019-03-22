# -*- coding: UTF-8 -*-
import cx_Oracle as cx
from docx import Document
from pandas import DataFrame

from src.gui.util.config import LOG
from .oracle import TAB_COMMENT_STATS, TABLE_DES_STATS, GET_ALL_USER_OBJ

__all__ = ['get_table_docs', 'write_to_execl', 'write_to_word', 'split_table_name']


def get_table_docs(cursor: cx.CURSOR, table_owner: str, table_name: str) -> dict:
    """
    获取单张表的表结构描述。
    :rtype: dict
    :param cursor: 数据库游标
    :param table_owner: 用户
    :param table_name: 表名
    :return: `dict` or None
    """
    table_name = table_name.upper().strip()
    if isinstance(cursor, cx.Cursor):
        _tab_comments = TAB_COMMENT_STATS % (repr(table_owner), repr(table_name))
        _tab_desc = TABLE_DES_STATS % (repr(table_owner), repr(table_name))
        _comment = cursor.execute(_tab_comments).fetchall()
        _desc = cursor.execute(_tab_desc).fetchall()
        comment = DataFrame(_comment)
        desc = DataFrame(_desc)
        return {'col_desc': comment, 'tab_comment': desc}
    else:
        LOG.error('get_table_docs: \'cursor\' is not a cx_Oracle.Cursor type. ')
        LOG.error('get_table_docs: additional info: %s ' % type(cursor))
    return None


def write_to_execl(writer, dd, style=None):
    """
    将Dict中的数据 写入execl文件的单个sheet页中
    todo: 待加入对样式的支持
    :param writer: execl操作对象
    :param dd: dict 数据
    :param style: SHEET页样式
    :return:
    """
    col_desc = DataFrame(dd.get('tab_comment'))
    tab_comments = DataFrame(dd.get('col_desc'))
    sheet_name: str = tab_comments.loc[0, 1]

    if not str(tab_comments.loc[0, 3]).strip().isspace() and tab_comments.loc[0, 3] is not None:
        sheet_name = str(tab_comments.loc[0, 3])
    print(sheet_name)
    print(type(sheet_name))
    try:
        tab_comments.to_excel(writer, sheet_name=sheet_name, index=False, header=['用户', '表名', '类型', '备注'], startrow=0)
        col_desc.to_excel(writer, sheet_name=sheet_name, index=False,
                          header=['序号', '列英文名', '数据类型', '是否可空', '默认值', '列中文名'], startrow=5, inf_rep="")
    except:
        LOG.error("%s write failed!" % sheet_name)
    finally:
        writer.save()


def write_to_word(document: Document, dd: dict, style=None, **kwargs):
    """
    写入word 文件
    todo: 待加入对样式的支持
    :param document: `Document` 对象
    :param dd: 内容，字典
    :param style 表格样式
    :return: none
    """
    table_comment = DataFrame(dd.get('tab_comment'))
    table_desc = DataFrame(dd.get('col_desc'))
    table_name: str = table_desc.loc[0, 1]
    heading_level = kwargs.get("heading_level")
    if heading_level is None:
        heading_level = 2
    # 写标题
    document.add_heading(table_name, level=heading_level)
    document.add_paragraph('\n')

    #  #############################################表描述#########################################
    table = document.add_table(rows=table_desc.shape[0] + 1, cols=4)
    hdr_cells = table.rows[0].cells
    #  写表头 '用户', '表名', '类型', '备注'
    hdr_cells[0].text = '用户'
    hdr_cells[1].text = '表名'
    hdr_cells[2].text = '类型'
    hdr_cells[3].text = '备注'
    #  写内容
    table_desc = table_desc.fillna("")
    print(table_desc)
    for i in range(table_desc.shape[0]):  # 获取数据框行数 shape[0]
        # row_cells = table.add_row().cells
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = table_desc.iloc[i, 0]
        row_cells[1].text = table_desc.iloc[i, 1]
        row_cells[2].text = table_desc.iloc[i, 2]
        row_cells[3].text = table_desc.iloc[i, 3]

    # 列字典
    document.add_paragraph('\n')
    table = document.add_table(rows=table_comment.shape[0] + 1, cols=6)
    hdr_cells = table.rows[0].cells
    #  写表头 '序号', '列英文名', '数据类型', '是否可空', '默认值', '列中文名'
    hdr_cells[0].text = '列英文名'
    hdr_cells[1].text = '数据类型'
    hdr_cells[2].text = '是否可空'
    hdr_cells[3].text = '默认值'
    hdr_cells[4].text = '列中文名'
    #  写内容
    table_comment = table_comment.fillna("")
    for i in range(table_comment.shape[0]):
        # row_cells = table.add_row().cells
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = str(table_comment.iloc[i, 0])
        row_cells[1].text = str(table_comment.iloc[i, 1])
        row_cells[2].text = str(table_comment.iloc[i, 2])
        row_cells[3].text = str(table_comment.iloc[i, 3])
        row_cells[4].text = str(table_comment.iloc[i, 4])
        row_cells[5].text = str(table_comment.iloc[i, 5])
    # 加分页符
    document.add_page_break()


def split_table_name(table_name: str):
    """
    分割表名
    :param table_name: 表名 “owner.table”
    :return: 返回 元组类型的表名（owner， table）
    """
    return table_name.split('.', maxsplit=2)


def get_all_table(cursor: cx.Cursor, owner: str):
    """
    获取某个用户下的所有表
    :param cursor: 游标 类型 cx_Oracle.Cursor
    :param owner:  用户 schema
    :return:
    """
    STMTS = GET_ALL_USER_OBJ % repr(owner.upper().strip())
    res = cursor.execute(STMTS).fetchall()
    tab_list = []
    for i, tables in enumerate(res):
        tab_list.append(tables[1])
    return tab_list
